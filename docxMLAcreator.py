from docx import Document
from docx.shared import Pt, Inches
from tkinter import Tk
from tkinter import filedialog
from datetime import date
import os

def setDirectory():
    directory = filedialog.askdirectory()
    root.destroy()
    return directory 

def createPath():
    file_path = setDirectory()
    file_name = input("Enter a name for your file:\n")
    file_path += '/%s.docx' % (file_name)
    return file_path

def setMLAstyle(document):
    style = document.styles['Normal']
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 2.0
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    return

def createFile(name,date,prof,course,title,file_path):
    document = Document()
    setMLAstyle(document)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    document.add_paragraph(name,style='Normal')
    document.add_paragraph(prof,style='Normal')
    document.add_paragraph(course,style='Normal')
    document.add_paragraph(date,style='Normal')
    titleP = document.add_paragraph(title,style='Normal')
    titleP.paragraph_format.alignment = 1
    document.save(file_path)
    return 

if __name__ == "__main__":
    name = "John Doe"
    root = Tk()
    today = date.today()
    date = today.strftime("%B %d, %Y")
    file_path = createPath() 
    
    prof = 'Professor ' + input("Enter name of your professor:\n")
    course = input("Enter course number and section number (Ex. CS-260-003, or however your school does it):\n")
    title = input("Enter title of your paper:\n")
    createFile(name,date,prof,course,title,file_path)
    os.startfile(file_path)
